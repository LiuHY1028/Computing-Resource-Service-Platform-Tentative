#!/usr/bin/env python3
"""Generate the complete product requirements document from approved product rules."""

from __future__ import annotations

import argparse
import os
import re
import sys
import tempfile
import uuid
import zipfile
from dataclasses import dataclass
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml import parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree


DOC_TITLE = "算力资源服务平台功能需求说明书"
VERSION = "V1.0"
CHAPTER_TITLES = [
    "产品概述",
    "建设目标、范围与实现边界",
    "用户角色与术语",
    "总体功能架构与跨模块关系",
    "资源商城",
    "统一购买与交易流程",
    "云服务器管理",
    "物理机管理",
    "软件中心",
    "存储购买与管理",
    "文件管理",
    "镜像管理",
    "网络与访问",
    "订单管理",
    "账单管理",
    "收银台",
    "操作记录",
    "价格与计费",
    "状态流转与数据模型",
    "校验、异常与权限边界",
    "功能验收标准",
]

FIGURES = [
    ("01-marketplace.png", "资源商城与资源筛选", 5),
    ("02-cloud-purchase.png", "云服务器购买配置", 6),
    ("03-physical-purchase.png", "物理机购买配置", 6),
    ("04-cloud-resources-actions.png", "云服务器列表与资源操作", 7),
    ("05-storage-purchase.png", "存储购买配置", 10),
    ("06-storage-management.png", "存储资源管理", 10),
    ("07-file-management.png", "文件管理工作区", 11),
    ("08-public-images.png", "公共镜像管理", 12),
    ("09-custom-images.png", "自定义镜像管理", 12),
    ("10-network-rules.png", "网络访问规则", 13),
    ("11-order-detail.png", "订单详情", 14),
    ("12-checkout.png", "收银台付款", 16),
    ("13-operation-records.png", "全局操作记录", 17),
]

INK = RGBColor(18, 32, 58)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(91, 104, 125)
LIGHT_FILL = "F2F4F7"
BLUE_FILL = "E8EEF5"
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
EMBEDDED_FONT_PATH = Path("/System/Library/Fonts/Supplemental/Arial Unicode.ttf")
EMBEDDED_FONT_NAME = "Arial Unicode MS"


@dataclass(frozen=True)
class ModuleSpec:
    prefix: str
    name: str
    roles: str
    preconditions: str
    entry: str
    capabilities: tuple[str, ...]
    main_flow: tuple[str, ...]
    alternative_flow: tuple[str, ...]
    states: str
    fields: tuple[str, ...]
    validations: tuple[str, ...]
    feedback: str
    exceptions: tuple[str, ...]
    relations: str
    acceptance: tuple[str, ...]


MODULES: dict[int, ModuleSpec] = {
    5: ModuleSpec(
        "MKT", "资源商城", "所有已登录用户", "内置商品目录、价目目录和可购状态已加载",
        "#/marketplace",
        ("按资源类型、站点、计算类型和计费模式筛选商品", "比较云服务器与物理机规格", "进入适配商品的配置流程"),
        ("进入资源商城", "选择云服务器或物理机", "筛选并核对规格与价格", "进入商品配置"),
        ("不可购规格保留说明但禁止进入配置", "筛选无结果时允许清空条件后重试"),
        "商品可购状态是唯一入口判断，不与资源运行状态混用。",
        ("商品编号", "资源类型", "站点", "核心规格", "计费模式", "起始价格", "可购状态"),
        ("不可购商品按钮必须禁用并说明原因", "价格只能读取统一价目目录", "商品类型决定后续配置字段"),
        "筛选即时生效；进入配置时保留商品编号。",
        ("目录损坏时显示错误与重试入口", "商品引用的站点或价格缺失时禁止购买"),
        "连接购买草稿、价格目录、镜像选择和订单确认。",
        ("用户可完成筛选与比较", "不可购商品不可绕过", "进入配置后商品与站点一致"),
    ),
    6: ModuleSpec(
        "PUR", "统一购买与交易流程", "资源购买用户、财务付款用户",
        "商品可购且配置数据通过校验", "商品配置页、确认订单页、收银台",
        ("统一配置、确认订单、支付三阶段", "保存可恢复的购买草稿", "创建冻结配置与价格的订单快照", "预付费创建账单并进入收银台", "按量订单直接进入履约"),
        ("选择商品并配置", "进入确认订单", "核对配置与费用", "创建订单及账单", "完成付款", "执行本地履约", "进入资源或订单详情"),
        ("返回已完成配置阶段修改草稿", "取消待支付订单后从快照重新配置", "支付失败后保留订单并允许重试"),
        "草稿：配置→确认；订单：待支付→支付中→已支付→履约中→已完成；账单：未支付→支付中→已支付。",
        ("草稿版本", "商品类型", "配置阶段", "配置快照", "价格快照", "订单编号", "账单编号"),
        ("未来阶段不可直接进入", "配置不完整时不能进入确认订单", "订单创建后快照不可修改", "支付前不得创建或变更最终资源"),
        "步骤状态、按钮可用性、错误原因和后续入口保持一致。",
        ("刷新后草稿无效时返回配置并说明原因", "支付失败不得更新账单为已支付", "履约失败保留订单关联以便追踪"),
        "统一连接云服务器、物理机、存储、软件、订单、账单和收银台。",
        ("三类资源共用步骤语义", "金额和快照一致", "浏览器前进后退不破坏阶段", "支付后资源与订单一致"),
    ),
    7: ModuleSpec(
        "RES", "云服务器管理", "资源管理员、资源使用者",
        "云服务器已存在且用户可查看", "#/console/resources/cloud-servers",
        ("查看唯一主状态、规格、系统、网络与计费信息", "启动、停止、重启和修改名称", "创建续费订单", "设置自动续费", "制作自定义镜像", "维护标签和项目", "检查依赖后释放资源"),
        ("筛选并选择资源", "从行内或更多操作进入专用流程", "核对影响或填写参数", "确认操作", "更新本地状态并写入全局操作记录"),
        ("存储、网络、监控和操作记录使用正式页面跳转", "收费续费进入收银台，免费操作直接完成"),
        "创建中、运行中、已停止、重启中、即将到期、已到期、释放中、已释放、异常。",
        ("资源编号", "名称", "项目", "标签", "站点", "规格", "镜像", "到期时间", "自动续费", "唯一主状态"),
        ("运行中资源不得直接释放", "存在外挂存储或未完成订单时阻断释放", "制作镜像名称不得重复", "续费付款前不得更新到期时间"),
        "专用流程明确说明影响；成功后列表、详情和全局记录同步。",
        ("资源状态不允许执行操作时显示禁用原因", "释放失败不得留下部分变更", "制作镜像失败记录失败原因"),
        "连接镜像、存储、网络、续费订单、账单与操作记录。",
        ("更多操作不包含重装系统或变更配置", "每项操作产生独立结果", "释放进入终态", "详情和列表状态一致"),
    ),
    8: ModuleSpec(
        "PHY", "物理机管理", "资源管理员、资源使用者",
        "物理机已存在且用户可查看", "#/console/resources/physical-machines",
        ("查看物理机规格、硬件信息和唯一主状态", "启动、关机、重启和修改名称", "创建续租订单", "维护标签和项目", "跳转存储、网络、监控和操作记录", "检查依赖后释放"),
        ("选择物理机", "进入专用续租或运维流程", "确认参数与影响", "完成付款或免费操作", "同步资源与记录"),
        ("维护中物理机只允许查看和记录类操作", "续租失败时保留原到期时间"),
        "准备中、运行中、已关机、重启中、维护中、即将到期、已到期、释放中、已释放、异常。",
        ("资源编号", "名称", "项目", "站点", "CPU", "内存", "加速卡", "到期时间", "硬件摘要", "唯一主状态"),
        ("不提供制作系统镜像", "不提供自动续租设置", "续租付款前不得更新使用期限", "硬件健康只能作为辅助信息"),
        "续租显示新期限和费用；免费操作即时写入全局记录。",
        ("维护状态阻止不兼容操作", "依赖检查失败时不得进入释放中"),
        "连接存储、网络、续租订单、账单、监控和操作记录。",
        ("云服务器专属操作不出现", "续租订单与账单一致", "期限只在履约完成后更新"),
    ),
    9: ModuleSpec(
        "SW", "软件中心", "所有已登录用户",
        "软件目录和兼容规则可用", "#/software",
        ("浏览收费与免费软件", "查看兼容资源和费用策略", "收费软件创建订单", "免费软件创建安装任务"),
        ("选择软件", "核对兼容性和目标资源", "确认价格或免费安装", "创建订单或任务", "跟踪安装结果"),
        ("无兼容资源时引导先创建资源", "收费软件支付失败时不创建安装任务"),
        "等待执行、执行中、已完成、失败、已取消。",
        ("软件编号", "名称", "版本", "兼容计算类型", "费用策略", "目标资源", "任务状态"),
        ("目标资源必须可用且架构兼容", "收费项只能读取价格目录", "每项安装任务只有一个主状态"),
        "订单、安装任务和资源入口分别提供后续操作。",
        ("任务失败显示可读原因和重试入口", "资源释放后禁止新安装"),
        "连接资源、订单、账单、安装任务和操作记录。",
        ("收费软件进入交易流程", "免费软件不创建账单", "任务状态唯一"),
    ),
    10: ModuleSpec(
        "STO", "存储购买与管理", "资源管理员、存储管理员",
        "存储商品、站点和价格目录已加载", "#/console/storage/purchase、#/console/storage",
        ("选择云硬盘或高性能共享存储", "配置站点、性能、容量、数量、周期和自动续费", "按存储单元维护挂载计划", "查看实时报价并创建订单", "管理挂载、卸载、扩容和续费"),
        ("选择存储产品", "配置站点与性能", "配置容量数量和周期", "按需设置挂载", "核对报价", "确认订单并支付", "履约创建存储和挂载关系"),
        ("云硬盘可逐块选择目标或稍后挂载", "共享存储可为多个同站点资源分别设置路径和读写模式", "无适用资源时允许先购买后挂载"),
        "创建中、可用、挂载中、已挂载、卸载中、扩容中、续费中、即将到期、已到期、释放中、异常。",
        ("存储编号", "产品类型", "站点", "性能规格", "容量", "数量", "周期", "自动续费", "挂载计划", "价格快照"),
        ("容量和数量必须处于商品范围", "挂载目标必须同站点", "挂载路径必须合法且不能冲突", "订单数量必须等于履约资源数量"),
        "报价随配置同步；免费挂载明确不增加费用。",
        ("挂载目标失效时允许取消该挂载并继续购买", "支付失败不得创建存储", "未初始化云硬盘不得进入文件管理"),
        "连接价格、订单、账单、资源存储页、文件管理和操作记录。",
        ("多盘挂载不复用同一关系", "共享存储支持多目标独立参数", "订单账单和履约一致"),
    ),
    11: ModuleSpec(
        "FIL", "文件管理", "存储使用者",
        "存储已挂载且完成初始化", "#/console/storage/:storageId/files",
        ("浏览目录与文件", "创建目录", "上传、下载、重命名和删除", "查看文件任务与容量变化"),
        ("选择可用存储", "进入文件管理", "定位目录", "执行文件操作", "查看反馈和任务结果"),
        ("上传冲突时选择覆盖、跳过或重命名", "任务失败时保留失败记录并允许重试"),
        "文件任务：等待执行→执行中→已完成、失败或已取消。",
        ("文件编号", "名称", "类型", "父目录", "大小", "修改时间", "任务类型", "任务状态"),
        ("名称不得为空或包含非法路径字符", "容量不足时阻止上传", "删除目录前确认影响", "每个任务只有一个主状态"),
        "操作后刷新目录、容量和任务中心。",
        ("存储卸载或释放时禁止写操作", "浏览器不支持文件能力时提供明确原因"),
        "连接存储容量、挂载状态、文件任务和操作记录。",
        ("未初始化存储不可进入", "容量变化与文件操作一致", "失败任务不冒充成功"),
    ),
    12: ModuleSpec(
        "IMG", "镜像管理", "资源管理员、镜像管理员",
        "资源和镜像 Store 已加载", "#/console/images",
        ("管理公共镜像与自定义镜像两类", "从云服务器系统盘制作自定义镜像", "读取本地镜像文件元数据并导入记录", "筛选兼容镜像", "查看来源与关联资源", "删除未受限的自定义镜像"),
        ("选择镜像类型", "搜索或筛选", "查看详情", "制作或导入自定义镜像", "跟踪任务状态", "使用可用镜像创建资源"),
        ("制作或导入失败时记录失败原因", "已被资源使用的镜像删除前提示影响", "公共镜像仅允许查看与使用"),
        "从资源制作：制作中→可用或失败；文件导入：导入中→可用或失败。",
        ("镜像编号", "名称", "类型", "操作系统", "版本", "架构", "适用计算", "来源", "启动方式", "文件元数据", "失败原因", "状态"),
        ("自定义镜像必须有合法来源", "文件扩展名限 qcow2、raw、img、vhd、vhdx", "文件必须非空且不超过 30 GiB", "制作中或导入中不可删除"),
        "成功或失败均提供来源、关联资源和操作记录入口。",
        ("公共镜像不可修改删除", "来源资源不存在时禁止创建制作任务", "重名时阻止提交"),
        "连接云服务器操作、购买镜像选择、资源履约和操作记录。",
        ("一级分类仅两类", "购买只展示兼容且可用镜像", "关联资源动态派生"),
    ),
    13: ModuleSpec(
        "NET", "网络与访问", "资源管理员、网络维护人员",
        "资源存在且可管理访问规则", "#/console/network-access",
        ("按资源、类型、站点和状态筛选", "查看内外网地址和规则摘要", "使用 SSH、RDP、HTTP、HTTPS 模板", "创建自定义端口规则", "编辑、启用、停用和删除规则"),
        ("选择资源", "查看网络摘要", "选择常用模板或自定义", "配置允许来源", "创建并启用规则", "在全局操作记录追踪变更"),
        ("当前 IP 无法离线识别时禁用该选项", "全部来源允许继续但必须确认风险", "规则可先停用后再删除"),
        "规则状态仅为已启用或已停用，执行结果写入操作记录。",
        ("规则编号", "规则名称", "资源编号", "协议", "访问端口", "来源类型", "来源值", "说明", "状态"),
        ("端口范围为 1–65535", "指定 IP 与 CIDR 格式必须有效", "资源与规则必须同属有效引用", "同一资源不得创建完全重复规则"),
        "免费网络操作立即更新并写入全局记录，不创建订单或账单。",
        ("全部来源显示风险提示", "损坏旧规则迁移失败时回退网络初始数据", "删除失败不得留下半更新记录"),
        "连接资源详情、购买网络配置和全局操作记录。",
        ("页面不暴露方向、优先级和双端口", "resourceId 查询参数正确筛选", "无局部操作记录表"),
    ),
    14: ModuleSpec(
        "ORD", "订单管理", "购买用户、财务用户、运营查看者",
        "订单 Store 已加载", "#/console/orders",
        ("查看订单列表与唯一主状态", "按类型和状态筛选", "查看配置与价格快照", "待支付订单去支付或取消", "已完成订单进入资源"),
        ("进入订单列表", "定位订单", "查看详情和时间线", "执行支付、取消或查看资源"),
        ("取消待支付订单时同步取消未支付账单", "支付失败订单可返回收银台重试"),
        "待支付、支付中、已支付、履约中、已完成、已取消、支付失败、退款中、已退款。",
        ("订单编号", "订单类型", "商品或资源", "行项目", "价格快照", "主状态", "创建时间", "付款时间", "完成时间"),
        ("每个订单只有一个主状态", "预付费交易必须有有效账单", "快照金额等于账单金额", "已完成订单必须存在有效履约结果"),
        "操作可用性由唯一订单状态决定。",
        ("孤立订单不得展示为可操作", "已支付订单不得直接取消", "状态迁移失败时保持原状态"),
        "连接账单、收银台、资源、价格快照和操作记录。",
        ("待支付可取消", "已支付账单关系正确", "列表与详情状态一致"),
    ),
    15: ModuleSpec(
        "BIL", "账单管理", "财务用户、购买用户",
        "账单和关联订单有效", "#/console/bills",
        ("查看账单列表和详情", "查看关联订单与资源", "核对费用明细、出账和到期时间", "从未支付账单进入收银台"),
        ("进入账单列表", "筛选账单", "查看详情", "核对订单与费用", "完成支付或查看支付记录"),
        ("已取消账单只读展示", "退款账单显示退款过程与原订单关系"),
        "未支付、支付中、已支付、已取消、退款中、已退款。",
        ("账单编号", "订单编号", "账单类型", "商品或资源", "金额", "出账时间", "到期时间", "支付时间", "主状态"),
        ("账单必须关联有效订单", "金额来自订单价格快照", "已支付订单的账单必须已支付", "已取消订单不得存在已支付账单"),
        "账单详情提供订单、资源和付款记录跳转。",
        ("订单关系缺失时阻止付款", "账单损坏时只回退账单域，不影响资源域"),
        "连接订单、收银台、资源和价格快照。",
        ("金额与订单一致", "状态唯一", "不存在孤立账单"),
    ),
    16: ModuleSpec(
        "CHK", "收银台", "具有付款权限的用户",
        "订单处于待支付或支付失败且关联未支付账单有效", "#/checkout/:orderId",
        ("展示订单编号、商品配置和费用明细", "选择账户余额、企业付款账户或在线支付", "确认支付", "取消待支付订单", "展示履约与完成入口"),
        ("进入收银台", "核对订单和账单", "选择支付方式", "确认支付", "订单进入支付中", "账单支付完成", "订单履约并进入完成页"),
        ("支付失败保留账单未支付并允许重试", "取消订单同步取消账单", "完成订单进入只读结果页"),
        "收银台阶段由订单主状态映射，不单独维护第二套状态。",
        ("订单编号", "账单编号", "配置摘要", "费用明细", "应付金额", "支付方式", "结果入口"),
        ("仅待支付或支付失败订单允许确认支付", "应付金额必须等于账单金额", "支付完成前不得执行履约", "完成后不得修改价格快照"),
        "支付、履约和完成分别给出明确结果与后续入口。",
        ("本地支付更新失败时回滚订单和账单状态", "外部付款能力未接入时不得生成第三方交易号"),
        "连接购买步骤、订单、账单、履约、资源和操作记录。",
        ("付款与履约顺序正确", "取消保持关系一致", "完成页入口有效"),
    ),
    17: ModuleSpec(
        "OPS", "操作记录", "资源管理员、运营查看者、审计查看者",
        "全局操作记录 Store 已加载", "#/console/operation-records",
        ("统一记录资源、存储、镜像、网络、订单和文件操作", "按模块、资源、状态和时间筛选", "跳转关联订单、账单或资源"),
        ("操作开始时创建记录", "执行状态变化时更新同一记录", "完成或失败时写入结果", "用户在全局页面筛选和追踪"),
        ("纯页面跳转不新增记录", "创建订单与完成履约可使用关联标识串联"),
        "等待执行、执行中、已完成、失败、已取消。",
        ("记录编号", "模块", "操作类型", "资源编号", "订单编号", "账单编号", "关联标识", "主状态", "时间", "结果说明"),
        ("每个操作只保留一个当前主状态", "关联编号必须存在或为空", "局部页面不得维护第二份记录表"),
        "失败记录包含可读原因；筛选参数可由关联模块带入。",
        ("记录写入失败不得阻断核心状态的原子回滚", "旧重复记录在迁移后不得继续展示"),
        "连接所有免费运维操作、交易过程和全局导航。",
        ("网络页无内嵌记录", "资源详情无重复事实源", "筛选参数正确"),
    ),
    18: ModuleSpec(
        "PRI", "价格与计费", "购买用户、财务用户、产品维护人员",
        "统一价目目录和计价函数可用", "商品配置、确认订单、订单、账单与续费流程",
        ("按商品、镜像、存储和周期计算价格", "区分预付费与按量计费", "冻结历史价格快照", "为续费、续租和收费存储操作生成账单"),
        ("加载商品价格", "根据配置计算明细", "确认订单时展示", "创建订单时冻结", "账单读取快照", "历史详情继续使用原快照"),
        ("免费操作费用为零且不生成账单", "按量订单完成后按账期形成后付费账单"),
        "价格目录状态不作为资源或订单主状态展示。",
        ("币种", "整数分金额", "计费单位", "周期", "数量", "单价", "行项目", "总额", "快照版本"),
        ("禁止页面写死正式价格", "所有金额使用整数分计算", "行项目合计必须等于总额", "历史快照不受目录变化影响"),
        "报价变化即时反映在配置与确认阶段，创建订单后冻结。",
        ("价格缺失时禁止创建订单", "快照校验失败时不得进入收银台", "按量账期缺失时禁止履约"),
        "连接商城、购买、续费、续租、存储、软件、订单和账单。",
        ("价格来源唯一", "订单账单金额一致", "免费操作无账单"),
    ),
}


def set_run_font(run, *, size=None, bold=None, color=None, latin="Arial Unicode MS", east_asia="Arial Unicode MS"):
    run.font.name = latin
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), latin)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = "left" if side == "start" else "right" if side == "end" else side
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    width_dxa = [round(width * 1440) for width in widths]
    width_dxa[-1] += CONTENT_WIDTH_DXA - sum(width_dxa)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    for index, grid_col in enumerate(table._tbl.tblGrid.gridCol_lst):
        grid_col.set(qn("w:w"), str(width_dxa[index]))
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index])
            tc_w = cell._tc.get_or_add_tcPr().get_or_add_tcW()
            tc_w.set(qn("w:w"), str(width_dxa[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tr_pr = row._tr.get_or_add_trPr()
            cant_split = tr_pr.find(qn("w:cantSplit"))
            if cant_split is None:
                tr_pr.append(OxmlElement("w:cantSplit"))
    first_tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = first_tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        first_tr_pr.append(OxmlElement("w:tblHeader"))


def style_table_text(table, header=True):
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.05
                for run in paragraph.runs:
                    set_run_font(run, size=9.2, bold=(header and row_index == 0), color=INK)
            if header and row_index == 0:
                set_cell_shading(cell, LIGHT_FILL)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, text in enumerate(headers):
        table.rows[0].cells[index].text = text
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
    set_table_geometry(table, widths)
    style_table_text(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_para(doc, text="", *, bold=False, italic=False, color=INK, size=10.5, align=None, after=6, style=None):
    paragraph = doc.add_paragraph(style=style)
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = 1.10
    if text:
        run = paragraph.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color)
        run.italic = italic
    return paragraph


def add_bullets(doc, items):
    values = tuple(items)
    for index, item in enumerate(values):
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing = 1.10
        paragraph.paragraph_format.keep_with_next = index < len(values) - 1
        run = paragraph.add_run(item)
        set_run_font(run, size=10.3, color=INK)


def add_steps(doc, items):
    rows = [(index + 1, item, "用户完成本步骤后进入下一步骤") for index, item in enumerate(items)]
    add_table(doc, ("步骤", "行为", "结果"), rows, (0.65, 3.1, 2.75))


def add_page_field(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((fld_char, instr, separate, text, end))
    set_run_font(run, size=9, color=MUTED)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial Unicode MS"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (11.5, DARK_BLUE, 8, 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[style_name]
        style.font.name = "Arial Unicode MS"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    doc.styles["Heading 1"].paragraph_format.page_break_before = True

    for list_name in ("List Bullet", "List Number"):
        style = doc.styles[list_name]
        style.font.name = "Arial Unicode MS"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        style.font.size = Pt(10.3)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.167

    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_p.paragraph_format.space_after = Pt(0)
    run = header_p.add_run(DOC_TITLE)
    set_run_font(run, size=8.5, bold=True, color=MUTED)

    footer = section.footer
    footer_table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    footer_table.autofit = False
    footer_table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    footer_table.columns[0].width = Inches(4.8)
    footer_table.columns[1].width = Inches(1.7)
    left = footer_table.cell(0, 0).paragraphs[0]
    left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left_run = left.add_run(f"{VERSION} · {date.today().isoformat()}")
    set_run_font(left_run, size=8.5, color=MUTED)
    right = footer_table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    prefix = right.add_run("第 ")
    set_run_font(prefix, size=8.5, color=MUTED)
    add_page_field(right)
    suffix = right.add_run(" 页")
    set_run_font(suffix, size=8.5, color=MUTED)

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def embed_font(docx_path: Path, font_path: Path, font_name: str):
    if not font_path.exists():
        raise RuntimeError(f"嵌入字体不存在：{font_path}")
    font_key = uuid.uuid4()
    key = font_key.bytes_le[::-1]
    font_data = bytearray(font_path.read_bytes())
    for index in range(min(32, len(font_data))):
        font_data[index] ^= key[index % 16]

    with zipfile.ZipFile(docx_path, "r") as source:
        entries = {name: source.read(name) for name in source.namelist()}

    font_table = parse_xml(entries["word/fontTable.xml"])
    font_node = next((node for node in font_table.findall(qn("w:font")) if node.get(qn("w:name")) == font_name), None)
    if font_node is None:
        font_node = OxmlElement("w:font")
        font_node.set(qn("w:name"), font_name)
        font_table.append(font_node)
    for old in font_node.findall(qn("w:embedRegular")):
        font_node.remove(old)
    embed = OxmlElement("w:embedRegular")
    embed.set(qn("r:id"), "rIdEmbeddedFont")
    embed.set(qn("w:fontKey"), "{" + str(font_key).upper() + "}")
    embed.set(qn("w:subsetted"), "false")
    font_node.append(embed)
    entries["word/fontTable.xml"] = etree.tostring(
        font_table, xml_declaration=True, encoding="UTF-8", standalone=True
    )

    rels_name = "word/_rels/fontTable.xml.rels"
    if rels_name in entries:
        rels = parse_xml(entries[rels_name])
    else:
        rels = parse_xml(
            b'<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"/>'
        )
    for relationship in list(rels):
        if relationship.get("Id") == "rIdEmbeddedFont":
            rels.remove(relationship)
    relationship = etree.Element(
        "{http://schemas.openxmlformats.org/package/2006/relationships}Relationship"
    )
    relationship.set("Id", "rIdEmbeddedFont")
    relationship.set(
        "Type",
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/font",
    )
    relationship.set("Target", "fonts/ArialUnicodeMS.odttf")
    rels.append(relationship)
    entries[rels_name] = etree.tostring(
        rels, xml_declaration=True, encoding="UTF-8", standalone=True
    )
    entries["word/fonts/ArialUnicodeMS.odttf"] = bytes(font_data)

    content_types = parse_xml(entries["[Content_Types].xml"])
    package_ns = "http://schemas.openxmlformats.org/package/2006/content-types"
    if not any(node.get("Extension") == "odttf" for node in content_types):
        default = etree.Element(f"{{{package_ns}}}Default")
        default.set("Extension", "odttf")
        default.set(
            "ContentType",
            "application/vnd.openxmlformats-officedocument.obfuscatedFont",
        )
        content_types.append(default)
    entries["[Content_Types].xml"] = etree.tostring(
        content_types, xml_declaration=True, encoding="UTF-8", standalone=True
    )

    settings = parse_xml(entries["word/settings.xml"])
    for tag in ("embedTrueTypeFonts", "saveSubsetFonts"):
        node = settings.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            settings.append(node)
        node.set(qn("w:val"), "true")
    entries["word/settings.xml"] = etree.tostring(
        settings, xml_declaration=True, encoding="UTF-8", standalone=True
    )

    fd, temp_name = tempfile.mkstemp(suffix=".docx", dir=str(docx_path.parent))
    os.close(fd)
    temp_path = Path(temp_name)
    try:
        with zipfile.ZipFile(temp_path, "w", zipfile.ZIP_DEFLATED) as target:
            for name, data in entries.items():
                target.writestr(name, data)
        os.replace(temp_path, docx_path)
    finally:
        if temp_path.exists():
            temp_path.unlink()


def add_cover(doc):
    add_para(doc, "PRODUCT REQUIREMENTS DOCUMENT", size=10, bold=True, color=BLUE, after=24)
    add_para(doc, DOC_TITLE, size=28, bold=True, color=INK, after=10)
    add_para(doc, "面向产品、研发、测试与项目实施的完整功能需求", size=14, color=DARK_BLUE, after=34)
    add_table(
        doc,
        ("文档属性", "内容"),
        (
            ("文档版本", VERSION),
            ("文档状态", "正式发布"),
            ("适用范围", "算力资源商城、购买交易、资源管理及关联控制台能力"),
            ("实现形态", "纯前端离线单文件应用，本地状态持久化"),
            ("发布日期", date.today().isoformat()),
        ),
        (1.45, 5.05),
    )
    add_para(doc, "修订记录", size=13, bold=True, color=BLUE, after=6)
    add_table(
        doc,
        ("版本", "日期", "修订范围", "状态"),
        ((VERSION, date.today().isoformat(), "形成全产品功能、流程、状态、数据和验收基线", "发布"),),
        (0.8, 1.2, 3.8, 0.7),
    )
    doc.add_page_break()
    add_para(doc, "目录", size=20, bold=True, color=INK, after=16)
    for index, title in enumerate(CHAPTER_TITLES, start=1):
        p = add_para(doc, f"{index}. {title}", size=10.5, color=DARK_BLUE, after=3)
        p.paragraph_format.left_indent = Inches(0.15)
    add_para(doc, "目录依据一级业务章节生成；各模块采用一致的需求结构，便于产品、研发和测试追踪。", color=MUTED, size=9, after=0)


def add_heading(doc, text, level):
    return doc.add_heading(text, level=level)


def add_requirements(doc, spec: ModuleSpec):
    rows = []
    source = list(spec.capabilities) + list(spec.validations) + list(spec.acceptance)
    for index, statement in enumerate(source, start=1):
        rows.append((f"{spec.prefix}-{index:03d}", statement, "必须"))
    add_table(doc, ("需求编号", "功能要求", "级别"), rows, (1.05, 4.75, 0.7))


def add_figure(doc, image_path: Path, caption: str, figure_number: int):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    inline_shape = run.add_picture(str(image_path), width=Inches(6.35))
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", caption)
    doc_pr.set("title", caption)
    caption_p = add_para(
        doc,
        f"图 {figure_number} {caption}",
        size=9.2,
        color=MUTED,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=8,
    )
    caption_p.paragraph_format.keep_with_next = False


def add_module_chapter(doc, chapter_number, spec: ModuleSpec, figures_by_chapter, screenshot_dir):
    add_heading(doc, f"{chapter_number}. {spec.name}", 1)
    add_heading(doc, "功能目标", 2)
    add_para(doc, f"{spec.name}用于支撑本产品的核心业务闭环，并与统一订单、账单、资源状态和全局操作记录保持一致。")
    add_heading(doc, "使用角色", 2)
    add_para(doc, spec.roles)
    add_heading(doc, "前置条件", 2)
    add_para(doc, spec.preconditions)
    add_heading(doc, "功能入口", 2)
    add_para(doc, spec.entry)
    add_heading(doc, "功能列表", 2)
    add_bullets(doc, spec.capabilities)
    add_heading(doc, "主流程", 2)
    add_steps(doc, spec.main_flow)
    add_heading(doc, "备选流程", 2)
    add_bullets(doc, spec.alternative_flow)
    add_heading(doc, "状态流转", 2)
    add_para(doc, spec.states)
    add_heading(doc, "字段说明", 2)
    add_table(
        doc,
        ("字段", "业务含义"),
        tuple((field, f"{spec.name}在创建、查询、校验或关联时使用的正式字段。") for field in spec.fields),
        (1.65, 4.85),
    )
    add_heading(doc, "校验规则", 2)
    add_bullets(doc, spec.validations)
    add_heading(doc, "操作反馈", 2)
    add_para(doc, spec.feedback)
    add_heading(doc, "异常处理", 2)
    add_bullets(doc, spec.exceptions)
    add_heading(doc, "跨模块关联", 2)
    add_para(doc, spec.relations)
    add_heading(doc, "需求清单", 2)
    add_requirements(doc, spec)
    add_heading(doc, "验收标准", 2)
    add_bullets(doc, spec.acceptance)
    for filename, caption, _ in figures_by_chapter.get(chapter_number, []):
        add_figure(doc, screenshot_dir / filename, caption, FIGURE_INDEX[(filename, caption)])


FIGURE_INDEX = {(filename, caption): index for index, (filename, caption, _) in enumerate(FIGURES, start=1)}


def add_chapters(doc, screenshot_dir):
    figures_by_chapter: dict[int, list[tuple[str, str, int]]] = {}
    for filename, caption, chapter in FIGURES:
        figures_by_chapter.setdefault(chapter, []).append((filename, caption, chapter))

    add_heading(doc, "1. 产品概述", 1)
    add_para(doc, "本平台提供计算资源、存储、镜像、基础网络访问、软件、订单、账单和操作记录的一体化自助服务。用户可从资源商城完成配置、确认订单、付款和履约，并在控制台管理资源全生命周期。")
    add_heading(doc, "1.1 产品定位", 2)
    add_bullets(doc, ("面向企业内部算力资源使用与管理", "以资源、交易和运维三类事实模型支撑跨模块一致性", "在离线单文件运行边界内提供可恢复的本地状态"))
    add_heading(doc, "1.2 核心价值", 2)
    add_table(doc, ("价值", "说明"), (
        ("统一交易", "购买、续费、续租和收费存储操作使用订单与账单"),
        ("统一状态", "同一对象在同一信息层级只展示一个主状态"),
        ("统一关系", "资源、存储、镜像、网络、订单和操作记录使用有效引用"),
        ("统一交付", "正式页面可在离线环境完成完整流程和状态校验"),
    ), (1.4, 5.1))
    add_heading(doc, "1.3 产品边界", 2)
    add_para(doc, "本版本不连接真实支付、镜像仓库、网络编排、资源交付、外部 IP 检测、密码或凭据服务。页面仅更新当前应用内的本地状态，不生成第三方交易号或外部基础设施结果。")

    add_heading(doc, "2. 建设目标、范围与实现边界", 1)
    add_heading(doc, "2.1 建设目标", 2)
    add_bullets(doc, ("建立一致的购买和交易链路", "让保留的资源操作形成独立且可验证的业务结果", "建立公共与自定义镜像两类模型", "提供普通用户可理解的基础访问规则", "形成可开发、可测试、可实施的完整需求基线"))
    add_heading(doc, "2.2 产品范围", 2)
    add_table(doc, ("范围", "包含能力"), (
        ("商品与购买", "资源商城、三类资源配置、确认订单、收银台"),
        ("资源与运维", "云服务器、物理机、存储、文件、镜像、网络"),
        ("交易与追踪", "订单、账单、操作记录、价格与计费"),
    ), (1.5, 5.0))
    add_heading(doc, "2.3 实现边界", 2)
    add_para(doc, "所有正式数据来自前端内置数据和本地持久化状态；应用不发起 HTTP 请求，不依赖服务器或同目录运行资产。")

    add_heading(doc, "3. 用户角色与术语", 1)
    add_heading(doc, "3.1 用户角色", 2)
    add_table(doc, ("角色", "主要职责"), (
        ("资源使用者", "购买、查看和使用已授权资源"),
        ("资源管理员", "管理生命周期、标签、项目、镜像和网络"),
        ("存储管理员", "购买、挂载、初始化和管理存储与文件"),
        ("财务用户", "核对订单、账单和付款状态"),
        ("运营查看者", "查看全局操作记录和关系状态"),
    ), (1.45, 5.05))
    add_heading(doc, "3.2 核心术语", 2)
    add_table(doc, ("术语", "定义"), (
        ("订单", "记录一次购买、续费、续租或收费存储行为及其冻结快照"),
        ("账单", "记录订单对应的应收、支付或退款事实"),
        ("履约", "付款后创建或更新应用内资源状态的过程"),
        ("公共镜像", "平台提供、用户只能查看和使用的镜像"),
        ("自定义镜像", "从云服务器制作或从本地镜像文件导入的用户镜像"),
        ("主状态", "对象在当前信息层级唯一展示的业务状态"),
    ), (1.45, 5.05))

    add_heading(doc, "4. 总体功能架构与跨模块关系", 1)
    add_heading(doc, "4.1 功能架构", 2)
    add_table(doc, ("层次", "模块", "主要关系"), (
        ("商品层", "资源商城、软件中心", "读取商品和价格并创建购买草稿"),
        ("交易层", "确认订单、订单、账单、收银台", "冻结快照并驱动付款与履约"),
        ("资源层", "云服务器、物理机、存储、镜像、网络", "保存唯一主状态与有效引用"),
        ("运维层", "文件管理、操作记录", "执行免费操作并提供统一追踪"),
    ), (1.0, 2.25, 3.25))
    add_heading(doc, "4.2 核心关系", 2)
    add_bullets(doc, (
        "预付费订单与账单一一关联，金额来自同一价格快照。",
        "资源记录商品、镜像和站点引用；镜像关联资源由资源的 imageId 动态派生。",
        "存储挂载关系按存储单元和目标资源分别保存。",
        "网络规则只引用有效资源，免费变更只写入全局操作记录。",
        "列表与详情从同一 Store 读取唯一主状态。",
    ))

    for chapter in range(5, 19):
        add_module_chapter(doc, chapter, MODULES[chapter], figures_by_chapter, screenshot_dir)

    add_heading(doc, "19. 状态流转与数据模型", 1)
    add_heading(doc, "19.1 单一主状态原则", 2)
    add_para(doc, "同一业务对象在同一信息层级只展示一个主状态。健康、风险、计费模式、自动续费、使用率、剩余时间和操作说明均使用辅助字段、指标或风险提示表达。")
    add_heading(doc, "19.2 状态集合", 2)
    add_table(doc, ("对象", "主状态集合"), (
        ("云服务器", "创建中、运行中、已停止、重启中、即将到期、已到期、释放中、已释放、异常"),
        ("物理机", "准备中、运行中、已关机、重启中、维护中、即将到期、已到期、释放中、已释放、异常"),
        ("存储", "创建中、可用、挂载中、已挂载、卸载中、扩容中、续费中、即将到期、已到期、释放中、异常"),
        ("订单", "待支付、支付中、已支付、履约中、已完成、已取消、支付失败、退款中、已退款"),
        ("账单", "未支付、支付中、已支付、已取消、退款中、已退款"),
        ("镜像任务", "制作中、导入中、可用、失败"),
        ("操作任务", "等待执行、执行中、已完成、失败、已取消"),
    ), (1.25, 5.25))
    add_heading(doc, "19.3 购买交易状态表", 2)
    add_steps(doc, ("配置商品", "确认订单", "创建订单和账单", "支付处理中", "支付成功", "履约中", "已交付"))
    add_heading(doc, "19.4 续费与续租状态表", 2)
    add_steps(doc, ("填写周期", "核对新期限与费用", "创建订单和账单", "支付", "履约", "更新到期时间"))
    add_heading(doc, "19.5 自定义镜像状态表", 2)
    add_table(doc, ("来源", "初始状态", "成功终态", "失败终态"), (
        ("云服务器系统盘", "制作中", "可用", "失败"),
        ("本地镜像文件", "导入中", "可用", "失败"),
    ), (1.8, 1.4, 1.4, 1.9))
    add_heading(doc, "19.6 资源释放状态表", 2)
    add_steps(doc, ("检查运行状态", "检查外挂存储", "检查未完成订单", "检查自动续费与网络规则", "确认资源名称", "进入释放中", "进入已释放"))
    add_heading(doc, "19.7 关键数据模型", 2)
    add_table(doc, ("模型", "关键字段", "关系约束"), (
        ("订单", "id、orderType、items、pricingSnapshot、status、时间", "预付费交易关联有效账单"),
        ("账单", "id、orderId、billType、amount、status、时间", "金额等于订单快照"),
        ("资源", "id、type、name、siteId、imageId、status、expiresAt", "状态唯一，引用有效"),
        ("镜像", "id、type、source、architecture、status、failureReason", "自定义镜像来源合法"),
        ("网络规则", "id、resourceId、protocol、port、sourceType、status", "免费操作且资源引用有效"),
        ("操作记录", "id、module、resourceId、relatedOrderId、status、时间", "全局唯一事实来源"),
    ), (1.1, 3.3, 2.1))
    add_heading(doc, "19.8 本地数据迁移", 2)
    add_bullets(doc, (
        "旧资源状态中的 resizing 迁移为 running，旧 releasing 迁移为 released。",
        "旧 resize 订单及其账单成对移除，其他价格快照与关系保留。",
        "旧 platform 镜像合并为 public，非法来源的待处理自定义镜像移除。",
        "旧网络规则按外部访问端口迁移为单端口，effective 映射为 enabled。",
        "损坏数据只回退相应领域的内置初始状态，不影响其他领域。",
    ))

    add_heading(doc, "20. 校验、异常与权限边界", 1)
    add_heading(doc, "20.1 通用校验", 2)
    add_bullets(doc, (
        "所有必填字段在进入下一阶段前完成校验。",
        "跨模块引用必须指向有效对象，孤立订单、账单、资源或规则不得进入正式操作。",
        "金额使用整数分计算并核对行项目合计。",
        "对象主状态必须属于对应状态集合。",
        "浏览器刷新和历史导航必须恢复合法阶段。",
    ))
    add_heading(doc, "20.2 异常处理", 2)
    add_table(doc, ("异常", "处理要求"), (
        ("本地持久化损坏", "仅重置对应领域并恢复内置初始数据"),
        ("价格或引用缺失", "阻止创建订单并指出缺失关系"),
        ("付款失败", "保持订单和账单可重试，不执行履约"),
        ("履约失败", "记录失败原因，不伪造已交付结果"),
        ("文件或镜像校验失败", "保留用户输入并说明具体规则"),
        ("网络规则风险", "提示全部来源风险但允许用户明确继续"),
    ), (1.6, 4.9))
    add_heading(doc, "20.3 权限边界", 2)
    add_para(doc, "当前版本提供角色职责和操作边界说明，但不接入外部身份、组织或权限服务。公共镜像只读；付款、资源释放、镜像删除等重要操作必须经过明确确认。")
    add_heading(doc, "20.4 当前实现边界", 2)
    add_para(doc, "支付、镜像制作、文件导入、网络规则和资源履约只更新本应用内状态；不连接真实第三方服务，不生成外部交易号、真实资源编排结果、外部 IP、密码或凭据。")

    add_heading(doc, "21. 功能验收标准", 1)
    add_heading(doc, "21.1 购买与交易", 2)
    add_bullets(doc, (
        "云服务器、物理机和存储使用一致的配置、确认订单和支付阶段。",
        "订单创建后配置与价格快照不可修改。",
        "预付费订单与账单金额一致，支付前不提前更新资源。",
        "取消和支付失败保持订单账单状态一致。",
    ))
    add_heading(doc, "21.2 资源与运维", 2)
    add_bullets(doc, (
        "云服务器和物理机操作入口符合产品边界，保留操作具有独立结果。",
        "续费、续租、自动续费、标签项目、镜像制作和释放均可验证。",
        "资源释放完成依赖检查并进入终态。",
        "列表与详情使用同一唯一主状态。",
    ))
    add_heading(doc, "21.3 存储、镜像与网络", 2)
    add_bullets(doc, (
        "多云硬盘逐块保存挂载计划，共享存储逐资源保存路径和读写模式。",
        "镜像一级分类仅公共和自定义，自定义来源和状态可追踪。",
        "基础网络规则支持模板、单端口、来源类型和启停。",
        "免费挂载、网络和镜像运维不创建账单。",
    ))
    add_heading(doc, "21.4 质量与交付", 2)
    add_bullets(doc, (
        "目标视口无整页横向溢出、菜单或弹窗裁切、页面假死或控制台错误。",
        "关系、价格、存储、状态、交易、镜像和网络验证全部通过。",
        "离线单文件可通过 file 协议直接打开且无外部运行时依赖。",
        "本说明书可由 Word 或 WPS 正常打开，章节、表格、图片和关系完整。",
    ))


FORBIDDEN_TEXT = (
    "Codex",
    "Prompt",
    "任务编号",
    "截图问题",
    "UI规范",
    "页面美化",
    "圆角",
    "阴影",
)


def verify_document(path: Path):
    document = Document(path)
    h1 = [p.text.strip() for p in document.paragraphs if p.style.name == "Heading 1"]
    text = "\n".join(p.text for p in document.paragraphs)
    for table in document.tables:
        text += "\n" + "\n".join(cell.text for row in table.rows for cell in row.cells)
    missing = [title for index, title in enumerate(CHAPTER_TITLES, start=1) if f"{index}. {title}" not in h1]
    forbidden = [token for token in FORBIDDEN_TEXT if token in text]
    if len(h1) != 21 or missing:
        raise RuntimeError(f"一级章节校验失败：count={len(h1)}, missing={missing}")
    if forbidden:
        raise RuntimeError(f"正文包含禁止表达：{forbidden}")
    if len(document.inline_shapes) < 13:
        raise RuntimeError(f"图片数量不足：{len(document.inline_shapes)}")
    if len(document.tables) < 45:
        raise RuntimeError(f"表格数量不足：{len(document.tables)}")
    ids = set(re.findall(r"\b(?:MKT|PUR|RES|PHY|SW|STO|FIL|IMG|NET|ORD|BIL|CHK|OPS|PRI)-\d{3}\b", text))
    if len(ids) < 80:
        raise RuntimeError(f"需求编号数量不足：{len(ids)}")
    with zipfile.ZipFile(path) as archive:
        names = set(archive.namelist())
        media = [name for name in names if name.startswith("word/media/")]
        rels = archive.read("word/_rels/document.xml.rels").decode("utf-8")
        if len(media) < 13 or "image" not in rels:
            raise RuntimeError("图片媒体关系不完整")
        if "[Content_Types].xml" not in names or "word/document.xml" not in names:
            raise RuntimeError("DOCX 核心关系缺失")
    if path.stat().st_size < 500_000:
        raise RuntimeError(f"文件大小异常：{path.stat().st_size}")
    return {
        "chapters": len(h1),
        "tables": len(document.tables),
        "images": len(document.inline_shapes),
        "requirements": len(ids),
        "bytes": path.stat().st_size,
    }


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--screenshots-dir", type=Path, required=True)
    parser.add_argument("--output", type=Path, required=True)
    args = parser.parse_args()

    missing_images = [filename for filename, _, _ in FIGURES if not (args.screenshots_dir / filename).exists()]
    if missing_images:
        raise SystemExit(f"缺少功能截图：{missing_images}")

    args.output.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    configure_document(document)
    add_cover(document)
    add_chapters(document, args.screenshots_dir)
    document.core_properties.title = DOC_TITLE
    document.core_properties.subject = "算力资源服务平台完整功能需求"
    document.core_properties.author = "算力资源服务平台项目组"
    document.core_properties.keywords = "算力资源, 订单, 账单, 存储, 镜像, 网络"
    document.core_properties.comments = "正式功能需求基线"
    document.save(args.output)
    embed_font(args.output, EMBEDDED_FONT_PATH, EMBEDDED_FONT_NAME)
    result = verify_document(args.output)
    print(
        "PRD generated: "
        + ", ".join(f"{key}={value}" for key, value in result.items()),
        flush=True,
    )


if __name__ == "__main__":
    try:
        main()
    except Exception as error:
        print(f"PRD generation failed: {error}", file=sys.stderr)
        raise
